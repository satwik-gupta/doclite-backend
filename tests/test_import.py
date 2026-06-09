"""Import tests: .md and .docx uploads preserve headings and lists."""
import io

from docx import Document as Docx

from tests.conftest import auth_header


def test_markdown_import_preserves_structure(client):
    header = auth_header(client, "alice")
    md = b"# Title\n\n## Section\n\n- one\n- two\n\n1. first\n2. second\n\n**bold**"
    res = client.post(
        "/api/documents/import",
        headers=header,
        files={"file": ("notes.md", md, "text/markdown")},
    )
    assert res.status_code == 201, res.text
    html = res.json()["content_html"]
    assert "<h1>" in html and "<h2>" in html
    assert "<ul>" in html and "<ol>" in html
    assert "<strong>" in html


def test_docx_import_preserves_structure(client):
    header = auth_header(client, "alice")
    d = Docx()
    d.add_heading("Doc Title", level=1)
    d.add_heading("Subsection", level=2)
    p = d.add_paragraph()
    p.add_run("important").bold = True
    d.add_paragraph("bullet a", style="List Bullet")
    d.add_paragraph("bullet b", style="List Bullet")
    d.add_paragraph("step 1", style="List Number")
    buf = io.BytesIO()
    d.save(buf)

    res = client.post(
        "/api/documents/import",
        headers=header,
        files={
            "file": (
                "report.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 201, res.text
    html = res.json()["content_html"]
    assert "<h1>" in html and "<h2>" in html
    assert "<ul>" in html and "<ol>" in html
    assert "<strong>" in html
    assert res.json()["title"] == "report"


def test_unsupported_type_rejected(client):
    header = auth_header(client, "alice")
    res = client.post(
        "/api/documents/import",
        headers=header,
        files={"file": ("x.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )
    assert res.status_code == 415
    assert res.json()["error"]["code"] == "unsupported_file_type"
